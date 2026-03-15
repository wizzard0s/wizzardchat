"""Generate a Word document with all integration-tabs test screenshots for review."""
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

SHOTS = Path("C:/Users/nico.debeer/WIZZARDCHAT/screenshots")
OUT   = Path("C:/Users/nico.debeer/WIZZARDCHAT/Integration_Tabs_Review.docx")

STEPS = [
    ("00_login.png",                        "Step 0 — Login",
     "Playwright logs in as admin and lands on the dashboard."),
    ("01_queues_list.png",                  "Step 1 — Queues list",
     "The Queues page loads with existing queue cards rendered."),
    ("01b_queues_debug.png",                "Step 1b — Queues debug",
     "Debug screenshot confirming card count before modal interaction."),
    ("02_queue_integrations_tab.png",       "Step 2 — Queue modal: Integrations tab",
     "The queue edit modal opens and the Integrations tab is activated."),
    ("03_queue_integration_slots_filled.png","Step 3 — Slots filled",
     "Slot 1 set to YouTube, slot 2 set to CNN News, Override Campaign checked."),
    ("04_queue_saved.png",                  "Step 4 — Queue saved",
     "Save button clicked; modal closes and the queue list reloads."),
    ("05_campaigns_list.png",               "Step 5 — Campaigns list",
     "The Campaigns page loads with campaign cards."),
    ("06_campaign_integration_slots_filled.png", "Step 6 — Campaign slot filled",
     "Campaign modal: slot 1 set to News24 and saved."),
    ("07_agent_panel_idle.png",             "Step 7 — Agent panel (idle)",
     "Agent panel loaded while no session is active — integration tab bar area visible."),
    ("08_agent_dom_verified.png",           "Step 8 — Agent DOM verified",
     "#integrationTabBar, #chatSlot, and #iSlot_1 through #iSlot_5 present in DOM."),
]

def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return h

def add_image_step(doc, png_path, title, caption):
    add_heading(doc, title, level=2)
    p = doc.add_paragraph(caption)
    p.runs[0].font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    p.runs[0].font.size = Pt(10)
    if png_path.exists():
        doc.add_picture(str(png_path), width=Inches(6.0))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph(f"[Screenshot not found: {png_path.name}]")
    doc.add_paragraph()  # spacer

doc = Document()

# ── Title ──────────────────────────────────────────────────────────────────
title = doc.add_heading("WizzardChat — Integration Tabs Feature: Test Review", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

meta = doc.add_paragraph(
    f"Generated: {datetime.date.today().isoformat()}    |    "
    "Branch: wizzardchat/main    |    Server: http://localhost:8092"
)
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.runs[0].font.size = Pt(9)
meta.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)
doc.add_paragraph()

# ── Summary ────────────────────────────────────────────────────────────────
add_heading(doc, "Summary", level=1)
doc.add_paragraph(
    "This document captures the Playwright end-to-end test run that verifies the "
    "Integration Tabs feature added to WizzardChat. The feature allows up to 5 "
    "configurable web-page URLs to be saved against a queue or campaign. Those URLs "
    "render as iFrame tabs in the agent panel whenever a session from that queue is active."
)
doc.add_paragraph()

# ── Scope ──────────────────────────────────────────────────────────────────
add_heading(doc, "Changes verified", level=1)
items = [
    "app/models.py — integration_urls JSONB column on Queue and Campaign",
    "app/schemas.py — integration_urls field on QueueCreate, QueueOut, CampaignCreate, CampaignOut",
    "templates/queues.html — Integrations tab (5 name/URL slots + Override Campaign checkbox)",
    "static/js/queues.js — _fillIntegrationSlots, _readIntegrationSlots helpers; saveQueue payload",
    "templates/campaigns.html — Integrations tab (5 name/URL slots)",
    "static/js/campaigns.js — same helpers; saveCampaign payload",
    "templates/agent.html — #integrationTabBar, #chatSlot, #iSlot_1..5 structure + CSS",
    "static/js/agent.js — _loadIntegrationTabs() called in openSession()",
    "DB migration — integration_urls JSONB added to chat.chat_queues and chat.chat_campaigns",
]
for item in items:
    p = doc.add_paragraph(item, style="List Bullet")
    p.runs[0].font.size = Pt(10)
doc.add_paragraph()

# ── Test pass/fail ─────────────────────────────────────────────────────────
add_heading(doc, "Test result", level=1)
doc.add_paragraph(
    "All assertions passed. Queue integration URLs (YouTube, CNN News) and campaign URL "
    "(News24) were saved and verified via API round-trip. "
    "Agent panel DOM elements confirmed present."
)
doc.add_paragraph()

# ── Screenshots ────────────────────────────────────────────────────────────
add_heading(doc, "Screenshots", level=1)

for fname, title, caption in STEPS:
    add_image_step(doc, SHOTS / fname, title, caption)

doc.save(str(OUT))
print(f"Saved: {OUT}")
