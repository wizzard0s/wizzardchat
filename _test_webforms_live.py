"""
Live webforms test — verifies that the agent panel renders the correct webform
tabs for two queue configurations:

  Campaign "Webform Demo Campaign" slots:  News24 | YouTube | Disney
  Queue A  (no override, linked to campaign): shows campaign slots
  Queue B  (override=true):                  shows CNN + YouTube only

Outputs:
  _test_webforms_live.docx  — screenshots + pass/fail summary

Run with:
  & "C:\\Users\\nico.debeer\\WIZZARDTEST\\.venv\\Scripts\\python.exe" _test_webforms_live.py
"""

import os, sys, time, uuid, json, pathlib, textwrap
import httpx, psycopg2
from playwright.sync_api import sync_playwright
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# ─── Config ─────────────────────────────────────────────────────────────────
BASE_URL  = "http://localhost:8092"
DB_DSN    = "postgresql://postgres:postgres@localhost:5432/wizzardfrw"
ADMIN_USER, ADMIN_PASS = "admin", "M@M@5t3r"
EVIDENCE  = pathlib.Path(__file__).parent / "evidence" / "webforms"
EVIDENCE.mkdir(parents=True, exist_ok=True)

CAMPAIGN_NAME = "Webform Demo Campaign"
QUEUE_A_NAME  = "Webform Queue A (Campaign tabs)"
QUEUE_B_NAME  = "Webform Queue B (Override tabs)"

CAMPAIGN_SLOTS = [
    {"name": "News24",  "url": "https://www.news24.com"},
    {"name": "YouTube", "url": "https://www.youtube.com"},
    {"name": "Disney",  "url": "https://www.disneyplus.com"},
]
QUEUEB_SLOTS = [
    {"name": "CNN",     "url": "https://www.cnn.com"},
    {"name": "YouTube", "url": "https://www.youtube.com"},
]

results: list[dict] = []

# ─── Helpers ────────────────────────────────────────────────────────────────
def log(msg: str):
    print(f"  {msg}", flush=True)

def get_token(client: httpx.Client) -> str:
    r = client.post("/api/v1/auth/login",
                    data={"username": ADMIN_USER, "password": ADMIN_PASS})
    r.raise_for_status()
    return r.json()["access_token"]

def api(client: httpx.Client, method: str, path: str, **kw):
    return client.request(method, path, **kw)

def find_or_create(client, list_path, create_path, name_field, body, key="name"):
    r = api(client, "GET", list_path)
    r.raise_for_status()
    items = r.json() if isinstance(r.json(), list) else r.json().get("items", [])
    existing = next((i for i in items if i.get(key) == body.get(name_field)), None)
    if existing:
        log(f"  Found existing: {existing[key]} ({existing['id']})")
        return existing
    r2 = api(client, "POST", create_path, json=body)
    if not r2.is_success:
        print(f"    ERROR creating {body.get(name_field)}: {r2.text}")
        r2.raise_for_status()
    created = r2.json()
    log(f"  Created: {created[key]} ({created['id']})")
    return created

def inject_session(conn, queue_id: str, campaign_id: str, agent_id: str,
                   label: str) -> str:
    """Insert a minimal Interaction row so it appears in the agent panel."""
    sid = str(uuid.uuid4())
    key = f"test-wf-{uuid.uuid4().hex[:8]}"
    name = f"Test Visitor ({label})"
    meta = json.dumps({"name": name, "phone": "+27821234567",
                       "contact_id": str(uuid.uuid4())})
    with conn.cursor() as cur:
        cur.execute("""
            INSERT INTO chat.chat_interactions
                (id, session_key, connector_id, queue_id, agent_id, status,
                 channel, visitor_metadata, message_log, segments,
                 created_at, last_activity_at)
            VALUES
                (%s, %s,
                 (SELECT id FROM chat.chat_connectors LIMIT 1),
                 %s, %s, 'with_agent',
                 'chat', %s::jsonb, '[]'::jsonb, '[]'::jsonb,
                 NOW(), NOW())
        """, (sid, key, queue_id, agent_id, meta))
    conn.commit()
    log(f"  Injected session {key} → queue {queue_id[:8]}…")
    return key

def screenshot(page, name: str) -> pathlib.Path:
    path = EVIDENCE / f"{name}.png"
    page.screenshot(path=str(path), full_page=False)
    return path

def check(condition: bool, msg: str) -> bool:
    status = "PASS" if condition else "FAIL"
    icon   = "✓" if condition else "✗"
    print(f"    {icon} {status}: {msg}")
    results.append({"status": status, "msg": msg})
    return condition

# ─── Docx builder ───────────────────────────────────────────────────────────
def build_docx(screenshots: dict[str, pathlib.Path]):
    doc = Document()
    doc.core_properties.author = "WizzardChat Automated Test"

    h = doc.add_heading("WizzardChat — Webform Tabs Live Test", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Test Scenario", 1)
    doc.add_paragraph(
        'Campaign "Webform Demo Campaign" defines three webform slots: '
        "News24, YouTube, and Disney+.\n\n"
        "Queue A (no override) inherits these slots from the campaign.\n"
        "Queue B (override_campaign=true) defines its own two slots: CNN and YouTube, "
        "replacing the campaign slots entirely."
    )

    doc.add_heading("Test Results", 1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Check"
    hdr[1].text = "Status"
    for r in results:
        row = table.add_row().cells
        row[0].text = r["msg"]
        row[1].text = r["status"]
        if r["status"] == "FAIL":
            for cell in row:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.color.rgb = RGBColor(0xC0, 0x30, 0x30)

    for title, path in screenshots.items():
        if path and path.exists():
            doc.add_heading(title, 2)
            doc.add_picture(str(path), width=Inches(6))

    out = pathlib.Path(__file__).parent / "_test_webforms_live.docx"
    doc.save(str(out))
    print(f"\n  Docx saved → {out}")
    return out

# ─── Main ────────────────────────────────────────────────────────────────────
def main():
    print("\n═══ WizzardChat Webform Tabs — Live Test ═══\n")
    screenshots: dict[str, pathlib.Path] = {}

    # ── 1. API setup ─────────────────────────────────────────────────────────
    print("Step 1: API setup")
    with httpx.Client(base_url=BASE_URL, timeout=15) as client:
        tok = get_token(client)
        client.headers["Authorization"] = f"Bearer {tok}"

        # Campaign
        log("Campaign:")
        campaign = find_or_create(
            client,
            list_path="/api/v1/campaigns",
            create_path="/api/v1/campaigns",
            name_field="name",
            body={
                "name":         CAMPAIGN_NAME,
                "description":  "Live test campaign for webform tab verification",
                "campaign_type":"outbound_voice",
                "webform_urls": {"slots": CAMPAIGN_SLOTS},
                "is_active":    True,
            },
        )
        camp_id = campaign["id"]

        # Update if slots are missing (campaign existed before rename)
        if not campaign.get("webform_urls", {}).get("slots"):
            r = api(client, "PUT", f"/api/v1/campaigns/{camp_id}",
                    json={"webform_urls": {"slots": CAMPAIGN_SLOTS}})
            r.raise_for_status()
            log("  Updated campaign webform_urls slots")

        # Queue A — no override
        log("Queue A:")
        queue_a = find_or_create(
            client,
            list_path="/api/v1/queues",
            create_path="/api/v1/queues",
            name_field="name",
            body={
                "name":         QUEUE_A_NAME,
                "channel":      "chat",
                "campaign_id":  camp_id,
                "webform_urls": {"slots": [], "override_campaign": False},
                "is_active":    True,
            },
        )
        qa_id = queue_a["id"]

        # Ensure campaign_id + override=False are set
        if str(queue_a.get("campaign_id") or "") != str(camp_id) or \
                queue_a.get("webform_urls", {}).get("override_campaign"):
            r = api(client, "PUT", f"/api/v1/queues/{qa_id}",
                    json={"campaign_id": camp_id,
                          "webform_urls": {"slots": [], "override_campaign": False}})
            r.raise_for_status()
            log("  Updated Queue A campaign_id + override=False")

        # Queue B — override=true with CNN + YouTube
        log("Queue B:")
        queue_b = find_or_create(
            client,
            list_path="/api/v1/queues",
            create_path="/api/v1/queues",
            name_field="name",
            body={
                "name":         QUEUE_B_NAME,
                "channel":      "chat",
                "campaign_id":  camp_id,
                "webform_urls": {"slots": QUEUEB_SLOTS, "override_campaign": True},
                "is_active":    True,
            },
        )
        qb_id = queue_b["id"]

        if not queue_b.get("webform_urls", {}).get("override_campaign") or \
                len(queue_b.get("webform_urls", {}).get("slots", [])) != 2:
            r = api(client, "PUT", f"/api/v1/queues/{qb_id}",
                    json={"campaign_id": camp_id,
                          "webform_urls": {"slots": QUEUEB_SLOTS,
                                           "override_campaign": True}})
            r.raise_for_status()
            log("  Updated Queue B slots + override=True")

        # Get admin user ID
        r = api(client, "GET", "/api/v1/auth/me")
        r.raise_for_status()
        admin_id = r.json()["id"]

    check(True, "Campaign, Queue A, Queue B created/verified via API")

    # ── 2. Inject sessions ───────────────────────────────────────────────────
    print("\nStep 2: Inject test sessions into DB")
    conn = psycopg2.connect(DB_DSN)
    try:
        key_a = inject_session(conn, qa_id, camp_id, admin_id, "QueueA")
        key_b = inject_session(conn, qb_id, camp_id, admin_id, "QueueB")
    finally:
        conn.close()
    check(True, "Sessions injected with status=with_agent")

    # ── 3. Playwright: agent panel ───────────────────────────────────────────
    print("\nStep 3: Playwright — agent panel webform tabs")
    with sync_playwright() as pw:
        browser = pw.chromium.launch(headless=False, slow_mo=300)
        ctx     = browser.new_context(viewport={"width": 1600, "height": 900})
        page    = ctx.new_page()

        # Login
        log("Logging in…")
        page.goto(f"{BASE_URL}/login")
        page.fill("#username", ADMIN_USER)
        page.fill("#password", ADMIN_PASS)
        page.click("button[type=submit]")
        # Wait for redirect, then navigate to agent panel
        time.sleep(2)
        page.goto(f"{BASE_URL}/agent")
        page.wait_for_load_state("networkidle", timeout=10_000)
        screenshots["Login page"] = screenshot(page, "01_login")

        # Wait for WebSocket to connect and sessions to appear
        log("Waiting for agent panel…")
        time.sleep(3)
        page.reload()
        time.sleep(2)

        # ── Queue A session ──────────────────────────────────────────────────
        log(f"Opening Queue A session ({key_a[:16]}…)")
        session_a_sel = f"[data-key='{key_a}'], [data-session-key='{key_a}']"

        # Try clicking the session in the session list
        try:
            page.wait_for_selector(session_a_sel, timeout=8_000)
            page.click(session_a_sel)
        except Exception:
            # Fallback: find session card with visitor name
            try:
                page.click("text=Test Visitor (QueueA)", timeout=5_000)
            except Exception:
                log("  Session A not visible in panel — injecting via JS openSession()")
                page.evaluate(f"window.openSession && window.openSession('{key_a}')")
        time.sleep(2)
        screenshots["Queue A — session opened"] = screenshot(page, "02_queueA_session")

        # Verify webform tabs render: News24, YouTube, Disney
        tab_bar = page.locator("#webformTabBar")
        tab_bar_visible = tab_bar.is_visible()
        check(tab_bar_visible, "Queue A: webform tab bar is visible")
        if tab_bar_visible:
            tabs_html = tab_bar.inner_text()
            for expected in ["News24", "YouTube", "Disney"]:
                check(expected in tabs_html,
                      f"Queue A: tab '{expected}' present")
        screenshots["Queue A — webform tabs"] = screenshot(page, "03_queueA_tabs")

        # ── Queue B session ──────────────────────────────────────────────────
        log(f"Opening Queue B session ({key_b[:16]}…)")
        session_b_sel = f"[data-key='{key_b}'], [data-session-key='{key_b}']"
        try:
            page.wait_for_selector(session_b_sel, timeout=8_000)
            page.click(session_b_sel)
        except Exception:
            try:
                page.click("text=Test Visitor (QueueB)", timeout=5_000)
            except Exception:
                log("  Session B not visible — injecting via JS openSession()")
                page.evaluate(f"window.openSession && window.openSession('{key_b}')")
        time.sleep(2)
        screenshots["Queue B — session opened"] = screenshot(page, "04_queueB_session")

        tab_bar_b_visible = page.locator("#webformTabBar").is_visible()
        check(tab_bar_b_visible, "Queue B: webform tab bar is visible")
        if tab_bar_b_visible:
            tabs_html_b = page.locator("#webformTabBar").inner_text()
            check("CNN"     in tabs_html_b, "Queue B: tab 'CNN' present (override)")
            check("YouTube" in tabs_html_b, "Queue B: tab 'YouTube' present (override)")
            check("News24" not in tabs_html_b and "Disney" not in tabs_html_b,
                  "Queue B: News24 and Disney absent (overridden by queue)")
        screenshots["Queue B — webform tabs"] = screenshot(page, "05_queueB_tabs")

        # Switch to a webform tab and screenshot the iframe
        try:
            page.locator("#webformTabBar button.wftab:not(.chat-tab)").first.click()
            time.sleep(1)
            screenshots["Queue B — first webform tab active"] = screenshot(page, "06_queueB_tab_active")
        except Exception:
            pass

        browser.close()

    passed = sum(1 for r in results if r["status"] == "PASS")
    failed = sum(1 for r in results if r["status"] == "FAIL")
    print(f"\n  Results: {passed} passed, {failed} failed")

    # ── 4. Docx ──────────────────────────────────────────────────────────────
    print("\nStep 4: Generating docx")
    build_docx(screenshots)

    if failed:
        print("\n  SOME TESTS FAILED — review docx for details")
        sys.exit(1)
    else:
        print("\n  ALL TESTS PASSED")


if __name__ == "__main__":
    main()
