"""
Real-browser recording of the Zyxel Router Installation flow.

Drives Playwright Chromium (headed) through the live WizzardChat widget,
takes a screenshot after every meaningful state change, and assembles a
Word document with the real browser captures.

Run from CHATDEV root:
    .venv/Scripts/python.exe docs/record_zyxel_flow.py
"""

import asyncio, os, re, time, textwrap, sys, io

# Force UTF-8 stdout so emoji print statements work on Windows cp1252 consoles
if hasattr(sys.stdout, "buffer"):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

from pathlib import Path
from datetime import datetime

from playwright.async_api import async_playwright, expect

# ── Config ────────────────────────────────────────────────────────────────────
BASE_URL   = "http://localhost:8092"
API_KEY    = "AWOuAYW1kieis50pp7OzeRt-xyDm4HeoFvTbxd6Sn8A"   # zyxel-demo
WIDGET_URL = f"{BASE_URL}/chat-preview?key={API_KEY}"

VISITOR_NAME  = "Alex"
ROUTER_MODEL  = "Zyxel VMG3625-T20A"

# Follow-up messages to send to the AI bot
AI_TURNS = [
    "I need to connect it via VDSL on a Telkom line. Where do I start?",
    "I've plugged in the phone cable but the WAN LED is red. What does that mean?",
    "I reseated the DSL splitter and now WAN is amber — then green! Internet light is green too.",
    "done",   # exit keyword → triggers exit edge → closing node
]

SCREENSHOT_DIR = Path(__file__).parent / "screenshots_real"
SCREENSHOT_DIR.mkdir(exist_ok=True)

# AI can take 30-90 s for first qwen3 response
AI_TIMEOUT_MS  = 120_000
MSG_TIMEOUT_MS = 30_000        # for non-AI bot messages

# ── Helpers ───────────────────────────────────────────────────────────────────
shot_index = 0

async def shot(page, label: str) -> Path:
    """Take a full-page screenshot and one clipped to just the widget panel."""
    global shot_index
    shot_index += 1
    slug  = re.sub(r"[^a-z0-9]+", "_", label.lower()).strip("_")
    fname = SCREENSHOT_DIR / f"{shot_index:02d}_{slug}.png"

    # scroll widget into view first
    await page.evaluate("document.getElementById('wc-panel')?.scrollIntoView()")

    # clip to the widget panel for a clean widget-only image
    panel = page.locator("#wc-panel")
    box   = await panel.bounding_box()
    if box:
        await page.screenshot(
            path=str(fname),
            clip={"x": box["x"] - 4, "y": box["y"] - 4,
                  "width": box["width"] + 8, "height": box["height"] + 8},
        )
    else:
        await page.screenshot(path=str(fname), full_page=False)

    print(f"  📷  [{shot_index:02d}] {label}")
    return fname


async def count_bot_msgs(page) -> int:
    return await page.evaluate(
        "document.querySelectorAll('.wc-bubble.wc-bot').length"
    )


async def wait_for_bot_msg(page, expected_count: int, timeout_ms: int = MSG_TIMEOUT_MS):
    """Wait until at least `expected_count` bot bubbles are visible."""
    await page.wait_for_function(
        f"document.querySelectorAll('.wc-bubble.wc-bot').length >= {expected_count}",
        timeout=timeout_ms,
    )
    # small settle delay so text has rendered
    await page.wait_for_timeout(600)


async def send_message(page, text: str):
    """Type and send a visitor message — waits for input to be enabled first."""
    inp = page.locator("#wc-input")
    # Wait until the input field is no longer disabled (AI may hold it while processing)
    await page.wait_for_function(
        "!document.getElementById('wc-input')?.disabled",
        timeout=AI_TIMEOUT_MS,
    )
    await inp.click()
    await inp.fill(text)
    await page.wait_for_timeout(200)
    await page.locator("#wc-send").click()
    await page.wait_for_timeout(400)


# ── Transcript capture ────────────────────────────────────────────────────────
transcript: list[dict] = []


async def capture_transcript(page):
    """Read all current bubbles from the DOM."""
    bubbles = await page.evaluate("""
        () => {
            const msgs = document.querySelectorAll('#wc-msgs .wc-bubble');
            return Array.from(msgs).map(b => ({
                role: b.classList.contains('wc-visitor') ? 'visitor' : 'bot',
                text: b.innerText.trim()
            }));
        }
    """)
    return bubbles


# ── Main recording flow ───────────────────────────────────────────────────────
async def run():
    screenshots: list[tuple[str, Path]] = []   # (label, path)

    async with async_playwright() as pw:
        browser = await pw.chromium.launch(
            headless=False,
            args=["--window-size=1200,820", "--window-position=100,50"],
        )
        ctx  = await browser.new_context(viewport={"width": 1200, "height": 820})
        page = await ctx.new_page()

        # ── 1. Load the preview page ─────────────────────────────────────────
        print("\n🌐  Opening chat preview page …")
        await page.goto(WIDGET_URL, wait_until="networkidle")
        await page.wait_for_timeout(1000)

        # Screenshot of the closed widget (launcher visible)
        p = await shot(page, "Chat widget launcher — page loaded")
        screenshots.append(("Widget launcher visible on preview page", p))

        # ── 2. Open the widget ────────────────────────────────────────────────
        print("🖱️   Clicking launcher …")
        await page.locator("#wc-launcher").click()
        await page.wait_for_selector("#wc-panel.wc-open", timeout=5000)
        await page.wait_for_timeout(800)

        # ── 3. Wait for greeting message ──────────────────────────────────────
        print("⏳  Waiting for greeting …")
        await wait_for_bot_msg(page, 1)
        p = await shot(page, "Bot greeting — widget opened")
        screenshots.append(("Bot greets visitor (flow starts, greeting node fires)", p))

        # ── 4. Capture name ───────────────────────────────────────────────────
        print(f"✍️   Sending name: {VISITOR_NAME}")
        bot_count_before = await count_bot_msgs(page)   # read BEFORE sending
        await send_message(page, VISITOR_NAME)
        p = await shot(page, f"Visitor types name — {VISITOR_NAME}")
        screenshots.append((f"Visitor enters name: '{VISITOR_NAME}'", p))

        # Wait for bot to ask about router model
        await wait_for_bot_msg(page, bot_count_before + 1)
        p = await shot(page, "Bot asks for router model")
        screenshots.append(("Bot asks 'What is your router model?' (message node)", p))

        # ── 5. Capture router model ───────────────────────────────────────────
        print(f"✍️   Sending router model: {ROUTER_MODEL}")
        bot_count_before = await count_bot_msgs(page)   # read BEFORE sending
        await send_message(page, ROUTER_MODEL)
        p = await shot(page, f"Visitor types router model — {ROUTER_MODEL}")
        screenshots.append((f"Visitor enters router model: '{ROUTER_MODEL}'", p))

        # ── 6. Wait for AI bot first response ────────────────────────────────
        print("🤖  Waiting for AI bot first response (qwen3:8b — may take 30-90 s) …")
        await wait_for_bot_msg(page, bot_count_before + 1, timeout_ms=AI_TIMEOUT_MS)
        p = await shot(page, "AI bot first response — Zyxel setup overview")
        screenshots.append(("AI bot (qwen3:8b) responds with Zyxel setup guidance", p))

        # ── 7. Multi-turn AI conversation ─────────────────────────────────────
        for i, msg in enumerate(AI_TURNS):
            is_exit = msg.lower() in ("done", "exit", "bye")
            timeout  = MSG_TIMEOUT_MS if is_exit else AI_TIMEOUT_MS

            print(f"✍️   Turn {i+1}: {msg!r}")
            bot_count_before = await count_bot_msgs(page)   # read BEFORE sending
            await send_message(page, msg)
            p = await shot(page, f"Visitor turn {i+1} — {msg[:40]}")
            screenshots.append((f"Visitor: '{msg}'", p))

            if is_exit:
                # After exit keyword, flow should move to closing message node
                print("🔑  Exit keyword sent — waiting for closing message …")
                try:
                    await wait_for_bot_msg(page, bot_count_before + 1, timeout_ms=15_000)
                except Exception:
                    pass  # closing may or may not produce a visible bot message
                await page.wait_for_timeout(1500)
                p = await shot(page, "Flow closing — exit keyword triggered")
                screenshots.append(("Exit keyword fires 'exit' edge → closing message node → end", p))
            else:
                await wait_for_bot_msg(page, bot_count_before + 1, timeout_ms=timeout)
                p = await shot(page, f"AI bot turn {i+1} response")
                screenshots.append((f"AI bot reply (turn {i+1})", p))

        # ── 8. Final state ────────────────────────────────────────────────────
        await page.wait_for_timeout(1000)
        p = await shot(page, "Final — conversation complete")
        screenshots.append(("Final state — flow reached End node", p))

        # Capture full transcript from DOM
        transcript = await capture_transcript(page)
        print(f"\n📝  Transcript: {len(transcript)} messages captured from DOM")

        await browser.close()

    # ── Build Word document ───────────────────────────────────────────────────
    print("\n📄  Building report …")
    build_docx(screenshots, transcript)


def build_docx(screenshots: list[tuple[str, Path]], transcript: list[dict]):
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def add_heading(doc, text, level=1):
        p = doc.add_heading(text, level=level)
        return p

    def add_colored_row(table, cells_data, bg_hex="F0F4FF"):
        row = table.add_row()
        for i, val in enumerate(cells_data):
            c = row.cells[i]
            c.text = str(val)
            c.paragraphs[0].runs[0].font.size = Pt(9.5)
            tc   = c._tc
            tcPr = tc.get_or_add_tcPr()
            shd  = OxmlElement("w:shd")
            shd.set(qn("w:fill"), bg_hex)
            shd.set(qn("w:val"), "clear")
            tcPr.append(shd)
        return row

    doc = Document()

    # ── Page margins ─────────────────────────────────────────────────────────
    for sec in doc.sections:
        sec.top_margin    = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)

    # ── Cover ─────────────────────────────────────────────────────────────────
    t = doc.add_heading("WizzardChat — Live Flow Recording", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph("Zyxel Router Installation — AI Bot Demo")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(14)
    sub.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x88)

    meta = doc.add_paragraph(
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n"
        f"Connector API Key: {API_KEY[:16]}…\n"
        f"AI Model: wizzardai://ollama/qwen3:8b\n"
        f"WizzardChat: {BASE_URL}"
    )
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.runs[0].font.size = Pt(9)
    meta.runs[0].font.color.rgb = RGBColor(0x77, 0x77, 0x77)
    doc.add_page_break()

    # ── Screenshots ───────────────────────────────────────────────────────────
    add_heading(doc, "1. Live Chat Widget Screenshots", 1)
    doc.add_paragraph(
        "Each image below is a real Playwright browser screenshot taken from the live "
        "WizzardChat widget at http://localhost:8092. No mock rendering was used."
    )

    for idx, (label, path) in enumerate(screenshots, 1):
        if not path.exists():
            continue
        add_heading(doc, f"Stage {idx}: {label}", 2)
        try:
            doc.add_picture(str(path), width=Inches(5.2))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            doc.add_paragraph(f"[Image unavailable: {e}]")
        doc.add_paragraph("")   # spacer

    doc.add_page_break()

    # ── Transcript ────────────────────────────────────────────────────────────
    add_heading(doc, "2. Full Conversation Transcript", 1)
    doc.add_paragraph(
        "Captured directly from the browser DOM after the flow completed."
    )

    table = doc.add_table(rows=1, cols=3)
    table.style = "Light List Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "#"
    hdr[1].text = "Role"
    hdr[2].text = "Message"
    for c in hdr:
        c.paragraphs[0].runs[0].font.bold = True

    for i, msg in enumerate(transcript, 1):
        role   = msg.get("role", "")
        text   = msg.get("text", "")
        bg     = "E8F0FE" if role == "bot" else "FFF9E6"
        add_colored_row(table, [str(i), role.upper(), text], bg_hex=bg)

    doc.add_page_break()

    # ── Flow summary ──────────────────────────────────────────────────────────
    add_heading(doc, "3. Flow Topology", 1)
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Light Shading Accent 2"
    hdr = tbl.rows[0].cells
    for c, h in zip(hdr, ["Node type", "Label", "Description", "UUID (8)"]):
        c.text = h
        c.paragraphs[0].runs[0].font.bold = True

    nodes = [
        ("start",   "Start",               "Entry point of the flow",                         "ee8703fe"),
        ("message", "Greeting",            "Bot welcomes visitor",                            "961c9561"),
        ("input",   "Capture name",        "Visitor enters their name",                       "30f9d938"),
        ("message", "Ask router model",    "Bot asks for the router model number",            "de0996a0"),
        ("input",   "Capture router model","Visitor types their router model",                "8ca9f35b"),
        ("ai_bot",  "Zyxel AI Support",    "Multi-turn qwen3:8b support chat (max 8 turns)", "1471e4ee"),
        ("message", "Closing",             "Bot confirms success and says goodbye",           "b2f2f7fa"),
        ("end",     "End",                 "Flow completes",                                  "c484c0de"),
    ]
    bgs = ["F0F4FF", "E8F8F0"]
    for i, row in enumerate(nodes):
        add_colored_row(tbl, row, bg_hex=bgs[i % 2])

    # ── Footer note ───────────────────────────────────────────────────────────
    doc.add_paragraph("")
    note = doc.add_paragraph(
        "All screenshots in this document were captured from a real Playwright-driven "
        "Chromium session against a live WizzardChat instance. The AI responses are "
        "generated by qwen3:8b running locally via Ollama."
    )
    note.runs[0].font.size = Pt(8.5)
    note.runs[0].font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    out = Path(__file__).parent / "zyxel_live_recording.docx"
    doc.save(str(out))
    print(f"\n✅  Report saved: {out}")
    print(f"    Screenshots: {len([p for _, p in screenshots if p.exists()])} images")
    print(f"    Transcript:  {len(transcript)} messages")
    return out


if __name__ == "__main__":
    asyncio.run(run())
